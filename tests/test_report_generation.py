from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Cm, Mm

from src.build_report_docx import DOCX_PATH, FIGURE_INDEX_PATH, PROJECT_ROOT, AUDIT_PATH, build_report


@pytest.fixture(scope="session")
def generated_report() -> Path:
    if not DOCX_PATH.exists():
        build_report()
    return DOCX_PATH


def parse_selected_figure_paths() -> list[Path]:
    text = FIGURE_INDEX_PATH.read_text(encoding="utf-8")
    paths: list[Path] = []
    in_figures = False
    for line in text.splitlines():
        if line.strip() == "## Gambar":
            in_figures = True
            continue
        if line.strip() == "## Tabel":
            break
        if in_figures and line.startswith("| Gambar"):
            parts = [part.strip() for part in line.strip().strip("|").split("|")]
            if len(parts) == 5:
                paths.append(PROJECT_ROOT / parts[2].strip("`"))
    return paths


def test_docx_created_and_non_empty(generated_report: Path) -> None:
    assert generated_report.exists()
    assert generated_report.stat().st_size > 0


def test_docx_has_required_sections_and_no_placeholder(generated_report: Path) -> None:
    doc = Document(str(generated_report))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    for heading in [
        "BAB I PENDAHULUAN",
        "BAB II TINJAUAN PUSTAKA",
        "BAB III METODOLOGI PENELITIAN",
        "BAB IV HASIL DAN PEMBAHASAN",
        "BAB V KESIMPULAN DAN SARAN",
        "LAMPIRAN",
        "DAFTAR PUSTAKA",
    ]:
        assert heading in text
    assert "PERLU VERIFIKASI" not in text
    assert "salinan UCI" not in text
    assert "One-Hot Encoding" in text


def test_selected_figures_exist_and_are_embedded(generated_report: Path) -> None:
    figure_paths = parse_selected_figure_paths()
    for path in figure_paths:
        assert path.exists(), f"Missing figure path {path}"

    with zipfile.ZipFile(generated_report) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
    assert len(media) == len(figure_paths)


def test_docx_layout_and_tables(generated_report: Path) -> None:
    doc = Document(str(generated_report))
    section = doc.sections[0]
    assert abs(section.page_width - Mm(210)) <= 500
    assert abs(section.page_height - Mm(297)) <= 500
    assert abs(section.left_margin - Cm(4)) <= 500
    assert abs(section.top_margin - Cm(4)) <= 500
    assert abs(section.right_margin - Cm(3)) <= 500
    assert abs(section.bottom_margin - Cm(3)) <= 500
    assert len(doc.tables) >= 10


def test_audit_exists(generated_report: Path) -> None:
    assert AUDIT_PATH.exists()
    audit_text = AUDIT_PATH.read_text(encoding="utf-8")
    assert "Gambar yang dimasukkan" in audit_text
    assert "Tabel yang dimasukkan" in audit_text
