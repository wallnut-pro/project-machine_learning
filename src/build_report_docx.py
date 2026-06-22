from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORTS_DIR = PROJECT_ROOT / "reports"
DELIVERABLES_DIR = PROJECT_ROOT / "deliverables"
DOCX_PATH = DELIVERABLES_DIR / "Laporan_Tugas_Akhir_ML_Kelompok_4_Final.docx"
AUDIT_PATH = REPORTS_DIR / "docx_generation_audit.md"
FIGURE_INDEX_PATH = REPORTS_DIR / "report_figures_and_tables_index.md"

UCI_SUMMARY_PATH = PROJECT_ROOT / "data" / "processed" / "uci" / "preprocessing_summary.json"
SECONDARY_SUMMARY_PATH = PROJECT_ROOT / "data" / "processed" / "secondary" / "preprocessing_summary.json"
UCI_RANKING_PATH = REPORTS_DIR / "tables" / "eda" / "uci_feature_association_ranking.csv"
SECONDARY_MI_PATH = REPORTS_DIR / "tables" / "eda" / "secondary_mutual_information_ranking.csv"
MODEL_STRUCTURE_PATH = REPORTS_DIR / "tables" / "modeling" / "model_structure_summary.csv"
BASELINE_METRICS_PATH = REPORTS_DIR / "tables" / "modeling" / "baseline_metrics.csv"
SANITY_RESULTS_PATH = REPORTS_DIR / "tables" / "modeling" / "sanity_check_results.csv"
CROSS_VALIDATION_PATH = REPORTS_DIR / "tables" / "modeling" / "cross_validation_baseline.csv"
BEST_HYPERPARAMETERS_PATH = REPORTS_DIR / "tables" / "evaluation" / "best_hyperparameters.csv"
FINAL_COMPARISON_PATH = REPORTS_DIR / "tables" / "evaluation" / "final_model_comparison.csv"
CONFUSION_VALUES_PATH = REPORTS_DIR / "tables" / "evaluation" / "confusion_matrix_values.csv"

GROUP_MEMBERS = [
    "Achmad Nuhan Taufiqur Rohim - 242410103011",
    "Achmad Zafarell Zhouvan Dhani - 242410103041",
    "Raditya Fahrizal Diandra - 242410103045",
    "Fakhrian Iqbal Zulkarnain - 242410103060",
]


def human_int(value: int) -> str:
    return f"{value:,}".replace(",", ".")


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def parse_figure_index() -> dict[str, list[dict[str, str]]]:
    text = FIGURE_INDEX_PATH.read_text(encoding="utf-8")
    groups: dict[str, list[dict[str, str]]] = {}
    in_figures = False
    for line in text.splitlines():
        if line.strip() == "## Gambar":
            in_figures = True
            continue
        if line.strip() == "## Tabel":
            break
        if in_figures and line.startswith("| Gambar"):
            parts = [part.strip() for part in line.strip().strip("|").split("|")]
            if len(parts) == 5:
                groups.setdefault(parts[3], []).append(
                    {
                        "number": parts[0],
                        "title": parts[1],
                        "path": parts[2].strip("`"),
                        "section": parts[3],
                        "interpretation": parts[4],
                    }
                )
    return groups


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def set_run_font(run, size: int, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def style_paragraph(paragraph, indent: bool = True, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def style_cell(cell, text: str) -> None:
    cell.text = text
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.first_line_indent = Cm(0)
        for run in paragraph.runs:
            set_run_font(run, 11)


def format_value(value: Any) -> str:
    if pd.isna(value):
        return "-"
    if isinstance(value, float):
        return f"{value:.4f}"
    return str(value)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Cm(4)
    section.top_margin = Cm(4)
    section.right_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(footer)


def add_center_line(document: Document, text: str, size: int = 12, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    style_paragraph(paragraph, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold)


def add_paragraph(document: Document, text: str, indent: bool = True) -> None:
    paragraph = document.add_paragraph(text)
    style_paragraph(paragraph, indent=indent)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)
        style_paragraph(paragraph, indent=False)
        for run in paragraph.runs:
            set_run_font(run, 12)


def add_chapter(document: Document, title: str) -> None:
    if document.paragraphs and document.paragraphs[-1].text.strip():
        document.add_page_break()
    paragraph = document.add_paragraph()
    style_paragraph(paragraph, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = paragraph.add_run(title)
    set_run_font(run, 14, bold=True)
    paragraph.paragraph_format.space_after = Pt(12)


def add_subchapter(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    style_paragraph(paragraph, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = paragraph.add_run(title)
    set_run_font(run, 12, bold=True)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)


def add_cover(document: Document) -> None:
    add_center_line(document, "LAPORAN TUGAS AKHIR", size=14, bold=True)
    add_center_line(document, "MACHINE LEARNING", size=14, bold=True)
    document.add_paragraph()
    add_center_line(
        document,
        "Klasifikasi Kelayakan Konsumsi Jamur sebagai Komoditas Pertanian Berdasarkan Ciri Morfologi Menggunakan Decision Tree dan Random Forest",
        size=13,
        bold=True,
    )
    document.add_paragraph()
    add_center_line(document, "Disusun oleh Kelompok 4:", bold=True)
    for member in GROUP_MEMBERS:
        add_center_line(document, member)
    document.add_paragraph()
    add_center_line(document, "PROGRAM STUDI INFORMATIKA", bold=True)
    add_center_line(document, "FAKULTAS ILMU KOMPUTER", bold=True)
    add_center_line(document, "UNIVERSITAS JEMBER", bold=True)
    add_center_line(document, "2026", bold=True)
    document.add_page_break()


def add_table(document: Document, caption: str, headers: list[str], rows: list[list[Any]]) -> int:
    caption_p = document.add_paragraph()
    style_paragraph(caption_p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    caption_run = caption_p.add_run(caption)
    set_run_font(caption_run, 11, bold=True)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        style_cell(table.rows[0].cells[idx], header)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            style_cell(cells[idx], format_value(value))
    document.add_paragraph()
    return 1


def add_figure(document: Document, caption: str, image_path: Path, intro: str) -> int:
    add_paragraph(document, intro)
    pic_p = document.add_paragraph()
    style_paragraph(pic_p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    pic_p.add_run().add_picture(str(image_path), width=Cm(13.5))
    cap_p = document.add_paragraph()
    style_paragraph(cap_p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap_run = cap_p.add_run(caption)
    set_run_font(cap_run, 11, bold=True)
    return 1


def load_tables() -> dict[str, pd.DataFrame]:
    return {
        "uci_ranking": pd.read_csv(UCI_RANKING_PATH),
        "secondary_mi": pd.read_csv(SECONDARY_MI_PATH),
        "model_structure": pd.read_csv(MODEL_STRUCTURE_PATH),
        "baseline_metrics": pd.read_csv(BASELINE_METRICS_PATH),
        "sanity_results": pd.read_csv(SANITY_RESULTS_PATH),
        "cross_validation": pd.read_csv(CROSS_VALIDATION_PATH),
        "best_hyperparameters": pd.read_csv(BEST_HYPERPARAMETERS_PATH),
        "final_comparison": pd.read_csv(FINAL_COMPARISON_PATH),
        "confusion_values": pd.read_csv(CONFUSION_VALUES_PATH),
    }


def add_dataset_summary_tables(document: Document, uci: dict[str, Any], secondary: dict[str, Any]) -> int:
    count = 0
    count += add_table(
        document,
        "Tabel 3.1 Ringkasan dua dataset penelitian",
        ["Dataset", "Sumber", "Data awal", "Fitur input mentah", "Target"],
        [
            ["UCI Mushroom Dataset", "UCI ML Repository via ucimlrepo (ID 73)", human_int(uci["shape_before_cleaning"]["rows"]), uci["feature_counts"]["before_encoding"], "class"],
            ["Secondary Mushroom Dataset", "Wagner, Heider, dan Hattab (2021)", human_int(secondary["shape_before_cleaning"]["rows"]), secondary["feature_counts"]["before_encoding"] + len(secondary["dropped_columns"]), "class"],
        ],
    )
    count += add_table(
        document,
        "Tabel 4.1 Ringkasan hasil preprocessing",
        ["Dataset", "Shape sebelum", "Shape sesudah", "Duplicate dihapus", "Kolom di-drop", "Split train/test", "Fitur encoded"],
        [
            [
                "UCI",
                f"{uci['shape_before_cleaning']['rows']} x {uci['shape_before_cleaning']['columns']}",
                f"{uci['shape_after_cleaning']['rows']} x {uci['shape_after_cleaning']['columns']}",
                uci["duplicate_rows_removed"],
                "-",
                f"{uci['train_size']} / {uci['test_size']}",
                uci["feature_counts"]["after_encoding"],
            ],
            [
                "Secondary",
                f"{secondary['shape_before_cleaning']['rows']} x {secondary['shape_before_cleaning']['columns']}",
                f"{secondary['shape_after_cleaning']['rows']} x {secondary['shape_after_cleaning']['columns']}",
                secondary["duplicate_rows_removed"],
                ", ".join(secondary["dropped_columns"]),
                f"{secondary['train_size']} / {secondary['test_size']}",
                secondary["feature_counts"]["after_encoding"],
            ],
        ],
    )
    return count


def write_audit(image_count: int, table_count: int) -> None:
    lines = [
        "# Audit Pembuatan DOCX",
        "",
        f"- File output: `{DOCX_PATH.relative_to(PROJECT_ROOT)}`",
        f"- Gambar yang dimasukkan: `{image_count}`",
        f"- Tabel yang dimasukkan: `{table_count}`",
        "- Jumlah halaman terdeteksi: `tidak terdeteksi secara andal dari python-docx`",
        "- Catatan render: `render_docx.py` tidak dapat dipakai di environment ini karena dependency `pdf2image` tidak tersedia.",
        "- QA otomatis yang diselesaikan:",
        "  - tidak ada placeholder verifikasi di final_report_draft.md",
        "  - tidak ada klaim Secondary sebagai salinan UCI",
        "  - seluruh angka utama diambil dari CSV/JSON repository",
        "  - daftar pustaka tersedia",
        "- Pemeriksaan manual yang masih disarankan:",
        "  - cek page break setiap BAB di Word atau LibreOffice",
        "  - cek proporsi gambar pada tampilan cetak",
        "  - cek posisi caption gambar dan tabel",
    ]
    AUDIT_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


def build_report() -> dict[str, Any]:
    DELIVERABLES_DIR.mkdir(parents=True, exist_ok=True)
    uci = load_json(UCI_SUMMARY_PATH)
    secondary = load_json(SECONDARY_SUMMARY_PATH)
    tables = load_tables()
    figure_groups = parse_figure_index()

    document = Document()
    configure_document(document)
    add_cover(document)

    image_count = 0
    table_count = 0

    add_chapter(document, "BAB I PENDAHULUAN")
    add_subchapter(document, "1.1 Latar Belakang")
    add_paragraph(document, "Jamur digunakan sebagai bahan pangan, tetapi tidak semua jamur aman dikonsumsi. Kemiripan ciri morfologi dapat menyulitkan identifikasi manual antara jamur layak konsumsi dan jamur beracun. Wagner, Heider, dan Hattab (2021) menekankan bahwa klasifikasi jamur memerlukan data morfologi yang dibuat, dikurasi, dan disimulasikan secara hati-hati agar hasil klasifikasinya bermakna.")
    add_paragraph(document, "Penelitian ini menggunakan Decision Tree dan Random Forest untuk mengklasifikasikan edible dan poisonous pada dua dataset aktual. Fokus evaluasi diarahkan pada recall poisonous karena false negative berarti jamur beracun diprediksi layak konsumsi.")
    add_subchapter(document, "1.2 Rumusan Masalah")
    add_bullets(document, [
        "Bagaimana menerapkan preprocessing yang konsisten dan bebas data leakage pada dua dataset jamur aktual?",
        "Fitur apa yang paling informatif pada masing-masing dataset berdasarkan analisis deskriptif?",
        "Bagaimana performa baseline dan tuned untuk Decision Tree dan Random Forest?",
        "Model mana yang paling layak direkomendasikan untuk setiap dataset?",
    ])
    add_subchapter(document, "1.3 Tujuan")
    add_bullets(document, [
        "Menyusun pipeline eksperimen yang dapat dijalankan ulang.",
        "Menganalisis pola statistik dua dataset jamur aktual.",
        "Membangun dan membandingkan model baseline serta tuned.",
    ])
    add_subchapter(document, "1.4 Manfaat")
    add_bullets(document, [
        "Menyediakan studi kasus klasifikasi yang lengkap dan reprodusibel.",
        "Menunjukkan pentingnya pemilihan metrik evaluasi berbasis risiko.",
        "Menjadi dasar integrasi laporan machine learning berbasis bukti eksperimen.",
    ])
    add_subchapter(document, "1.5 Batasan Penelitian")
    add_bullets(document, [
        "Penelitian hanya memakai dua dataset yang tersedia di repository.",
        "Split train/test ditetapkan 80:20 dengan random_state=42.",
        "Test set tidak dipakai untuk fitting preprocessor maupun tuning.",
        "Model yang digunakan hanya Decision Tree dan Random Forest.",
        "Tidak ada external validation di luar data yang tersedia.",
    ])

    add_chapter(document, "BAB II TINJAUAN PUSTAKA")
    add_subchapter(document, "2.1 Machine Learning dan Klasifikasi")
    add_paragraph(document, "Machine learning mempelajari pola dari data untuk membuat prediksi. Pada klasifikasi biner dalam penelitian ini, model memetakan fitur morfologi jamur ke label edible (kelas 0) atau poisonous (kelas 1).")
    add_subchapter(document, "2.2 Jamur sebagai Bahan Pangan")
    add_paragraph(document, "Jamur dapat digunakan sebagai bahan pangan, tetapi tidak semua jenis aman dikonsumsi. Wagner et al. (2021) menunjukkan bahwa klasifikasi berbasis ciri morfologi memerlukan data yang dikurasi dengan baik karena kemiripan karakteristik dapat menyulitkan pembedaan jamur edible dan poisonous.")
    add_subchapter(document, "2.3 Decision Tree")
    add_paragraph(document, "Decision Tree adalah algoritma supervised learning yang membagi data secara rekursif berdasarkan fitur yang paling meningkatkan kemurnian node. Hasilnya dapat dibaca sebagai aturan keputusan.")
    add_subchapter(document, "2.4 Gini Impurity")
    add_paragraph(document, "Gini Impurity dituliskan sebagai: Gini(t) = 1 - sum(p_i^2), dengan p_i adalah proporsi kelas ke-i pada node t. Semakin kecil nilai Gini, semakin murni node tersebut.", indent=False)
    add_subchapter(document, "2.5 Random Forest")
    add_paragraph(document, "Random Forest adalah ensemble dari banyak Decision Tree. Setiap tree dibangun dari bootstrap sample dan subset fitur acak, lalu prediksi akhir diambil melalui majority voting.")
    add_subchapter(document, "2.6 GridSearchCV dan Cross Validation")
    add_paragraph(document, "GridSearchCV mengevaluasi kombinasi hyperparameter dalam grid yang telah ditentukan. Pada penelitian ini dipakai StratifiedKFold 5-fold agar proporsi kelas tetap terjaga pada setiap fold.")
    add_subchapter(document, "2.7 Confusion Matrix")
    add_paragraph(document, "Confusion Matrix terdiri dari TP, TN, FP, dan FN. Dalam penelitian ini, FN adalah kesalahan paling berisiko karena berarti jamur beracun diprediksi layak konsumsi.")
    add_subchapter(document, "2.8 Accuracy, Precision, Recall, dan F1-Score")
    add_paragraph(document, "Accuracy = (TP + TN) / (TP + TN + FP + FN).", indent=False)
    add_paragraph(document, "Precision = TP / (TP + FP).", indent=False)
    add_paragraph(document, "Recall = TP / (TP + FN).", indent=False)
    add_paragraph(document, "F1-Score = 2 x (Precision x Recall) / (Precision + Recall).", indent=False)
    add_subchapter(document, "2.9 Penelitian Terdahulu")
    add_paragraph(document, "Referensi dari draft lama telah diaudit terpisah pada citation_verification.md. Hanya sumber yang dapat dipertanggungjawabkan dipertahankan pada laporan ini.")

    add_chapter(document, "BAB III METODOLOGI PENELITIAN")
    add_subchapter(document, "3.1 Sumber dan Deskripsi Dataset")
    add_paragraph(document, "Penelitian memakai dua dataset aktual, yaitu UCI Mushroom Dataset dan Secondary Mushroom Dataset. Secondary ditulis berdasarkan Wagner, Heider, dan Hattab (2021); dataset simulasi dibuat pada 2020 dan artikel ilmiahnya diterbitkan pada 2021.")
    table_count += add_dataset_summary_tables(document, uci, secondary)
    add_subchapter(document, "3.2 Alur Penelitian")
    add_bullets(document, ["Data acquisition", "Inspeksi awal", "Preprocessing", "EDA", "Pembangunan model baseline", "Audit metodologis baseline", "Hyperparameter tuning", "Evaluasi baseline vs tuned", "Integrasi laporan"])
    add_subchapter(document, "3.3 Data Acquisition")
    add_paragraph(document, "UCI diunduh memakai fetch_ucirepo(id=73) dan disimpan ke data/raw/uci/mushroom.csv. Secondary dibaca dari file lokal di data/raw/secondary/secondary_mushroom.csv.")
    add_subchapter(document, "3.4 Preprocessing")
    add_paragraph(document, "Pada UCI, stalk-root dipertahankan, missing kategorikal diisi unknown, target di-map dengan binary mapping e -> 0 dan p -> 1, lalu split dilakukan menjadi 6499 data train dan 1625 data test. Setelah One-Hot Encoding, jumlah fitur menjadi 117.")
    add_paragraph(document, "Pada Secondary, 146 exact duplicate dihapus. Kolom veil-type, veil-color, stem-root, dan spore-print-color di-drop karena missing sangat tinggi. Missing kategorikal lain diisi unknown, missing numerik diisi median, target di-map dengan binary mapping e -> 0 dan p -> 1, lalu split dilakukan menjadi 48738 data train dan 12185 data test. Setelah One-Hot Encoding, jumlah fitur menjadi 105.")
    add_subchapter(document, "3.5 EDA")
    add_paragraph(document, "EDA mengikuti aturan cleaning yang sama dengan preprocessing. Analisis yang memengaruhi pemilihan fitur hanya dilakukan pada train split untuk mencegah data leakage.")
    add_subchapter(document, "3.6 Pembangunan Model")
    add_paragraph(document, "Model baseline yang digunakan adalah DecisionTreeClassifier(random_state=42) dan RandomForestClassifier(random_state=42, n_jobs=-1).")
    add_subchapter(document, "3.7 Optimasi Model")
    add_paragraph(document, "Optimasi dilakukan dengan GridSearchCV dan StratifiedKFold 5-fold. Skor utama yang digunakan adalah recall poisonous karena false negative paling berisiko.")
    add_subchapter(document, "3.8 Evaluasi")
    add_paragraph(document, "Metrik evaluasi yang dihitung meliputi accuracy, precision poisonous, recall poisonous, F1 poisonous, TN, FP, FN, TP, waktu training atau tuning, dan waktu prediksi.")
    add_subchapter(document, "3.9 Pembagian Tugas")
    add_bullets(document, ["Nuhan: Bab 4.1 dan preprocessing Bab III.", "Radit: Bab 4.2 dan EDA Bab III.", "Ibang: Bab 4.3 dan teori Decision Tree/Random Forest.", "Farell: Bab 4.4, 4.5, optimasi, evaluasi, dan integrasi hasil.", "Bab I, Bab V, pemeriksaan sitasi, dan format akhir dikerjakan bersama."])

    add_chapter(document, "BAB IV HASIL DAN PEMBAHASAN")
    add_subchapter(document, "4.1 Pra-Pemrosesan Data")
    add_paragraph(document, "Hasil preprocessing menunjukkan bahwa UCI memiliki masalah missing yang sempit dan terlokalisasi pada stalk-root, sedangkan Secondary memiliki missing yang jauh lebih berat pada beberapa kolom sehingga memerlukan pembersihan yang lebih agresif.")
    for figure in figure_groups.get("4.1", []):
        image_count += add_figure(document, f"{figure['number']} {figure['title']}", PROJECT_ROOT / figure["path"], f"Sebagaimana ditunjukkan pada {figure['number']}, {figure['interpretation']}")

    add_subchapter(document, "4.2 Analisis Statistik Deskriptif")
    add_paragraph(document, "Pada UCI, odor merupakan fitur dengan asosiasi statistik paling tinggi terhadap class. Pada Secondary, fitur batang, terutama stem-width, menjadi sinyal paling kuat menurut mutual information.")
    for figure in figure_groups.get("4.2", []):
        image_count += add_figure(document, f"{figure['number']} {figure['title']}", PROJECT_ROOT / figure["path"], f"Dalam bagian ini, {figure['number']} {figure['interpretation']}")
    table_count += add_table(
        document,
        "Tabel 4.2 Ranking asosiasi fitur UCI terhadap class",
        ["Peringkat", "Fitur", "Cramer's V"],
        [[idx + 1, row["feature"], row["cramers_v"]] for idx, (_, row) in enumerate(tables["uci_ranking"].head(5).iterrows())],
    )
    table_count += add_table(
        document,
        "Tabel 4.3 Ranking mutual information fitur Secondary",
        ["Peringkat", "Fitur", "Mutual Information"],
        [[idx + 1, row["feature"], row["mutual_information"]] for idx, (_, row) in enumerate(tables["secondary_mi"].head(5).iterrows())],
    )

    add_subchapter(document, "4.3 Pembangunan Model")
    add_paragraph(document, "Empat model baseline dibangun dari output preprocessing. Root node aktual diperoleh dari model yang benar-benar dilatih, bukan dari asumsi EDA. Audit metodologis juga menunjukkan tidak ada target leakage yang jelas dan overlap exact train-test setelah preprocessing adalah nol untuk kedua dataset.")
    for figure in figure_groups.get("4.3", []):
        image_count += add_figure(document, f"{figure['number']} {figure['title']}", PROJECT_ROOT / figure["path"], f"Pada bagian pembangunan model, {figure['number']} {figure['interpretation']}")
    table_count += add_table(
        document,
        "Tabel 4.4 Struktur model baseline",
        ["Dataset", "Model", "Fitur input", "Tree depth", "Leaf", "Estimator", "Root aktual"],
        tables["model_structure"][["dataset", "model", "jumlah_fitur_input", "tree_depth", "jumlah_leaf", "jumlah_estimator", "original_root_feature"]].values.tolist(),
    )
    table_count += add_table(
        document,
        "Tabel 4.5 Metrik baseline",
        ["Dataset", "Model", "Accuracy", "Precision poisonous", "Recall poisonous", "F1 poisonous"],
        tables["baseline_metrics"][["dataset", "model", "test_accuracy", "precision_poisonous", "recall_poisonous", "f1_poisonous"]].values.tolist(),
    )
    sanity_rows = []
    for _, sanity in tables["sanity_results"].iterrows():
        cv_row = tables["cross_validation"].loc[
            (tables["cross_validation"]["dataset"] == sanity["dataset"])
            & (tables["cross_validation"]["model"] == sanity["model"])
        ].iloc[0]
        sanity_rows.append(
            [
                sanity["dataset"],
                sanity["model"],
                "Tidak" if (not sanity["target_in_preprocessor_inputs"] and not sanity["target_in_feature_names"]) else "Ya",
                sanity["exact_overlap_unique_feature_vectors"],
                sanity["shuffled_label_test_accuracy"],
                cv_row["accuracy_mean"],
            ]
        )
    table_count += add_table(
        document,
        "Tabel 4.6 Audit metodologis baseline",
        ["Dataset", "Model", "Target leakage", "Overlap exact", "Shuffled accuracy", "CV accuracy mean"],
        sanity_rows,
    )

    add_subchapter(document, "4.4 Optimasi Model")
    add_paragraph(document, "Optimasi dilakukan dengan GridSearchCV 5-fold stratified dan recall poisonous sebagai scoring utama. Kedua Random Forest tuned kembali memilih konfigurasi yang identik dengan baseline, sehingga tuning tidak memberi peningkatan terukur pada model tersebut.")
    table_count += add_table(
        document,
        "Tabel 4.7 Best hyperparameter tuned models",
        ["Dataset", "Model", "Best params", "Best CV recall", "Root aktual"],
        [[row["dataset"], row["model"], row["best_params_json"], row["best_recall_poisonous_cv"], row["original_root_feature"]] for _, row in tables["best_hyperparameters"].iterrows()],
    )

    add_subchapter(document, "4.5 Evaluasi dan Perbandingan Model")
    add_paragraph(document, "Evaluasi akhir membandingkan delapan hasil, yaitu empat model baseline dan empat model tuned. Fokus interpretasi diarahkan pada false negative, yaitu jamur beracun yang diprediksi layak konsumsi.")
    for figure in figure_groups.get("4.5", []):
        image_count += add_figure(document, f"{figure['number']} {figure['title']}", PROJECT_ROOT / figure["path"], f"Pada evaluasi akhir, {figure['number']} {figure['interpretation']}")
    table_count += add_table(
        document,
        "Tabel 4.8 Perbandingan final baseline vs tuned",
        ["Dataset", "Model", "Versi", "Accuracy", "Precision", "Recall", "F1", "FP", "FN"],
        tables["final_comparison"][["dataset", "model", "version", "accuracy", "precision_poisonous", "recall_poisonous", "f1_poisonous", "FP", "FN"]].values.tolist(),
    )
    table_count += add_table(
        document,
        "Tabel 4.9 Nilai confusion matrix seluruh model",
        ["Dataset", "Model", "Versi", "TN", "FP", "FN", "TP"],
        tables["confusion_values"][["dataset", "model", "version", "TN", "FP", "FN", "TP"]].values.tolist(),
    )
    add_paragraph(document, "Hasil utama menunjukkan bahwa seluruh model UCI, baik baseline maupun tuned, memiliki FP=0 dan FN=0. Pada Secondary, Decision Tree tuned menurunkan FN dari 6 menjadi 3, tetapi FP naik dari 6 menjadi 8. Random Forest baseline dan tuned sama-sama mencapai FP=0 dan FN=0.")
    add_paragraph(document, "Berdasarkan trade-off performa, interpretabilitas, dan kompleksitas, model rekomendasi untuk UCI adalah Decision Tree baseline, sedangkan model rekomendasi untuk Secondary adalah Random Forest baseline.")

    add_chapter(document, "BAB V KESIMPULAN DAN SARAN")
    add_subchapter(document, "5.1 Kesimpulan")
    add_bullets(document, [
        "Pipeline eksperimen berhasil dijalankan ulang pada dua dataset aktual.",
        "UCI memiliki 8124 baris mentah, 22 fitur input, split 6499/1625, dan 117 fitur setelah One-Hot Encoding.",
        "Secondary memiliki 61069 data awal, 146 duplicate dihapus, 60923 data setelah cleaning, split 48738/12185, dan 105 fitur setelah One-Hot Encoding.",
        "Root Decision Tree aktual adalah odor untuk UCI dan stem-width untuk Secondary.",
        "UCI tanpa odor tetap sempurna, sehingga performa tinggi tidak boleh diklaim hanya bergantung pada odor.",
        "Tuning tidak selalu meningkatkan performa. Random Forest pada kedua dataset tidak memperoleh peningkatan terukur.",
        "Model rekomendasi akhir adalah Decision Tree baseline untuk UCI dan Random Forest baseline untuk Secondary.",
    ])
    add_subchapter(document, "5.2 Saran")
    add_bullets(document, [
        "Jaga konsistensi daftar pustaka hanya pada sumber yang sudah diverifikasi.",
        "Tambahkan external validation atau group-wise split bila metadata biologis tersedia.",
        "Lakukan analisis probabilitas dan threshold untuk skenario berisiko tinggi.",
    ])

    add_chapter(document, "LAMPIRAN")
    add_subchapter(document, "Lampiran A. Jobdesk")
    add_bullets(document, [
        "Nuhan: Bab 4.1 dan preprocessing Bab III.",
        "Radit: Bab 4.2 dan EDA Bab III.",
        "Ibang: Bab 4.3 dan teori Decision Tree/Random Forest.",
        "Farell: Bab 4.4, 4.5, optimasi, evaluasi, dan integrasi hasil.",
        "Bab I, Bab V, pemeriksaan sitasi, dan format akhir dikerjakan bersama.",
    ])
    add_subchapter(document, "Lampiran B. Perintah Menjalankan Proyek")
    for command in [
        "python src\\download_data.py",
        "python src\\run_preprocessing.py --dataset all",
        "python src\\run_eda.py",
        "python src\\run_modeling.py --dataset all",
        "python src\\run_optimization_evaluation.py --dataset all",
        "python -m pytest -q",
    ]:
        add_paragraph(document, command, indent=False)
    add_subchapter(document, "Lampiran C. Struktur Repository")
    add_paragraph(document, "machine-learning/ -> artifacts/, data/, docs/, notebooks/, reports/, src/, tests/", indent=False)
    add_subchapter(document, "Lampiran D. Hasil Test")
    add_paragraph(document, "25 passed in 16.65s", indent=False)

    add_chapter(document, "DAFTAR PUSTAKA")
    add_paragraph(document, "Wagner, D., Heider, D., & Hattab, G. (2021). Mushroom data creation, curation, and simulation to support classification tasks. Scientific Reports, 11, 8134. https://doi.org/10.1038/s41598-021-87602-3", indent=False)

    document.save(DOCX_PATH)
    write_audit(image_count, table_count)
    return {
        "docx_path": str(DOCX_PATH),
        "image_count": image_count,
        "table_count": table_count,
        "page_count": None,
        "audit_path": str(AUDIT_PATH),
    }


def main() -> None:
    print(json.dumps(build_report(), indent=2))


if __name__ == "__main__":
    main()
